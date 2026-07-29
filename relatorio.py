# -*- coding: utf-8 -*-
"""
Montagem da Justificativa Técnica Preliminar (PPR 2026).

Reproduz a estrutura de 11 seções do modelo aprovado. O que a carteira já sabe
sobre o imóvel entra preenchido e com a fonte declarada; o que depende de
análise entra como rascunho editável, para o técnico confirmar ou reescrever.
"""
from __future__ import annotations

import io
import re
from datetime import date

import pandas as pd

from nucleo import (
    AZUL,
    VERDE,
    br_data,
    br_num,
    parametros_do_imovel,
    txt,
)

MESES = [
    "janeiro", "fevereiro", "março", "abril", "maio", "junho",
    "julho", "agosto", "setembro", "outubro", "novembro", "dezembro",
]

# Seções do documento. 'auto' indica que a seção carrega um quadro montado a
# partir dos dados; o texto do técnico entra depois do quadro.
SECOES = [
    ("finalidade", "1", "Finalidade e metodologia", False,
     "Objetivo da análise, fontes consultadas e limites do documento."),
    ("identificacao", "2", "Identificação do imóvel", True,
     "Observações sobre convergência ou divergência entre os sistemas."),
    ("insercao", "3", "Formação e inserção urbanística", False,
     "Projeto urbanístico de origem, aprovação, inserção no PDOT e relação com o entorno."),
    ("uso", "4", "Uso e capacidade construtiva", True,
     "Leitura dos parâmetros: portes viáveis, usos compatíveis, condições de acesso."),
    ("fisica", "5", "Situação física e possessória", False,
     "O que a vistoria e as imagens indicam, e o que permanece sem informação conclusiva."),
    ("condicionantes", "6", "Condicionantes físicas e interferências", False,
     "Redes, servidões, faixas non aedificandi, topografia, ocupações, custos estimados."),
    ("ambiental", "7", "Aspectos ambientais, licenciamento e infraestrutura", False,
     "Processo ambiental, licenças, condicionantes e estágio da infraestrutura."),
    ("mercado", "8", "Contexto imobiliário e evidências de mercado", True,
     "Leitura dos comparáveis: resultados de certames, ágios, demanda e oferta concorrente."),
    ("avaliacao", "9", "Avaliação preliminar do ativo", True,
     "Ajuste a lista de atributos e condicionantes gerada a partir dos dados."),
    ("critica", "10", "Análise crítica", False,
     "Juízo do analista sobre o potencial do ativo e o peso relativo de cada condicionante."),
    ("encaminhamentos", "11", "Propostas de encaminhamento", True,
     "Texto complementar aos encaminhamentos marcados acima."),
]

ENCAMINHAMENTOS_PADRAO = [
    "Realizar estudo mercadológico, de geomarketing e de vocação imobiliária, com avaliação "
    "da demanda, da oferta concorrente e do perfil dos potenciais compradores e investidores.",
    "Avaliar cenários de comercialização, abrangendo venda integral, eventual reconfiguração "
    "da área, subdivisão urbanística ou comercialização por etapas.",
    "Atualizar a situação das contratações de infraestrutura que atendem ao parcelamento.",
    "Acompanhar os resultados das licitações dos imóveis vizinhos, registrando valores mínimos, "
    "propostas, ágios, número de concorrentes e usos pretendidos.",
    "Investigar potencialidades não contempladas nos usos ou parâmetros atuais, de forma a "
    "subsidiar, quando pertinente, proposta de alteração na próxima revisão da LUOS.",
]


def data_extenso(d: date | None = None) -> str:
    d = d or date.today()
    return f"{d.day} de {MESES[d.month - 1]} de {d.year}"


# ----------------------------------------------------------------------------
# Quadros montados a partir dos dados
# ----------------------------------------------------------------------------
def tabela_identificacao(im) -> list[list[str]]:
    if pd.notna(im.get("TB55_CARTORIO")):
        registro = f"{txt(im['TB55_CARTORIO'])} — registro de {br_data(im['TB55_DT_REGISTRO'])}"
    else:
        registro = "Sem registro cartorial vinculado na TB55"
    return [
        ["Cadastro imobiliário", txt(im["CD_IMOVEL"])],
        ["Endereço cartorial", txt(im["ENDERECO"])],
        ["Região administrativa", txt(im["REGADMIN"])],
        ["Área do terreno (polígono cadastral)", br_num(im["AREA_TERRENO"], 2, sufixo=" m²")],
        ["Proprietária", "Terracap"],
        ["Condição no GIU", txt(im["CD_DS_COND"])],
        ["Situação do lote (vigente)", f"{txt(im['SITUACAO_VIGENTE'])} — fonte: {txt(im['SITUACAO_FONTE'])}"],
        ["Projeto urbanístico", txt(im["IMU_PLAN_L"])],
        ["Norma de uso e ocupação", txt(im["PLANTA_DEC"])],
        ["Unidade de uso e ocupação", txt(im["PPR_Destinacao"]) if pd.notna(im.get("PPR_Destinacao"))
         else txt(str(im.get("DESTINACAO", ""))[:40])],
        ["Forma e posição do lote", f"{txt(im['FORMA_DESC'])} · {txt(im['PPR_Posiçao do imov'] if pd.notna(im.get('PPR_Posiçao do imov')) else im.get('POSICAO'))}"],
        ["Registro cartorial", registro],
        ["Avaliação mais recente", br_num(im["TB55_VL_AVALIACAO_RECENTE"], 2, prefixo="R$ ")],
        ["Situação comercial", "Registra alienação no histórico da TB55"
         if (im.get("TB55_QT_ALIENACOES") or 0) > 0
         else "Não foi identificada alienação anterior no histórico da TB55"],
    ]


def tabela_parametros(par: dict) -> list[list[str]]:
    pct = lambda v: br_num(v * 100, 0, sufixo="%") if v is not None and pd.notna(v) else "—"
    return [
        ["Área do terreno", br_num(par["area_terreno"], 2, sufixo=" m²")],
        ["Coeficiente básico de aproveitamento", br_num(par["cfa_basico"], 2)],
        ["Coeficiente máximo de aproveitamento", br_num(par["cfa_maximo"], 2)],
        ["Área máxima de construção", br_num(par["area_maxima"], 2, sufixo=" m²")],
        ["Taxa máxima de ocupação", pct(par["tx_ocup"])],
        ["Projeção máxima das edificações", br_num(par["projecao_maxima"], 2, sufixo=" m²")],
        ["Permeabilidade mínima", pct(par["tx_perm"])],
        ["Área permeável mínima", br_num(par["permeavel_minima"], 2, sufixo=" m²")],
        ["Altura máxima", br_num(par["alt_max"], 2, sufixo=" m") if par["alt_max"] is not None
         and pd.notna(par["alt_max"]) else "—"],
    ]


def usos_admitidos(im) -> str:
    """Extrai a descrição de usos do campo DESTINACAO da base cartográfica."""
    d = str(im.get("PPR_Detalhamento_Destinação") or "").strip()
    if not d or d.lower() == "nan":
        bruto = str(im.get("DESTINACAO") or "")
        partes = bruto.split(" - ", 1)
        d = partes[1] if len(partes) > 1 else bruto
    d = re.sub(r"\(VIDE ANEXO.*", "", d, flags=re.I).strip(" .,")
    d = " ".join(d.split())
    return d


def vizinhos(im, cons: pd.DataFrame, limite=12) -> pd.DataFrame:
    """Imóveis comparáveis da própria carteira: mesmo projeto urbanístico, ou,
    na falta dele, mesma região administrativa."""
    proj = im.get("IMU_PLAN_L")
    base = cons[cons["CD_IMOVEL"] != im["CD_IMOVEL"]]
    v = base[base["IMU_PLAN_L"] == proj] if pd.notna(proj) else base.iloc[0:0]
    criterio = f"mesmo projeto urbanístico ({txt(proj)})"
    if len(v) < 3:
        v = base[base["REGADMIN"] == im["REGADMIN"]]
        criterio = f"mesma região administrativa ({txt(im['REGADMIN'])})"
    v = v.sort_values("AREA_TERRENO", ascending=False).head(limite)
    return v, criterio


def tabela_vizinhos(v: pd.DataFrame) -> list[list[str]]:
    linhas = [["Cadastro", "Endereço", "Terreno (m²)", "Situação", "Avaliação (R$)"]]
    for r in v.itertuples():
        linhas.append([
            txt(r.CD_IMOVEL),
            txt(r.ENDERECO, 46),
            br_num(r.AREA_TERRENO),
            txt(r.SITUACAO_VIGENTE),
            br_num(r.TB55_VL_AVALIACAO_RECENTE),
        ])
    return linhas


# ----------------------------------------------------------------------------
# Rascunhos automáticos
# ----------------------------------------------------------------------------
def rascunhos(im, cons, obs, ali, par) -> dict:
    """Texto inicial de cada seção, derivado exclusivamente do que está na base."""
    r = {}
    cond = txt(im["CD_DS_COND"])
    r["finalidade"] = (
        f"A presente Justificativa Técnica Preliminar tem por finalidade registrar a análise do "
        f"imóvel selecionado no Sistema de Gestão de Imóveis Urbanos — GIU, na condição {cond}, "
        f"identificar seu potencial para o desenvolvimento de novos negócios e consolidar as "
        f"informações necessárias à definição de encaminhamentos posteriores.\n\n"
        f"A análise abrange os aspectos cadastrais, cartoriais, urbanísticos, ambientais, físicos, "
        f"jurídicos, de infraestrutura e de mercado, tendo como fontes os sistemas corporativos da "
        f"Terracap, os projetos urbanísticos aprovados, os processos administrativos, as imagens "
        f"aéreas e os documentos relacionados ao imóvel e ao seu entorno.\n\n"
        f"O presente documento corresponde à etapa preliminar de caracterização e prospecção. Não "
        f"se confunde com estudo de vocação imobiliária, estudo urbanístico preliminar, plano de "
        f"massas ou modelagem econômico-financeira, que constituem etapas posteriores de eventual "
        f"estruturação do negócio, conforme a metodologia prevista no Manual de Estudo Urbanístico "
        f"da DINEG."
    )

    ressalva_area = (
        "A área indicada é a do polígono cadastral e pode divergir da área registrada em matrícula; "
        "recomenda-se conferi-la na certidão antes da tramitação."
    )
    if im["SITUACAO_DIVERGE"]:
        r["identificacao"] = (
            f"Os sistemas divergem quanto à situação do lote: a base cartográfica registra "
            f"{txt(im['SITUACAO'])} e a vistoria da triagem PPR apurou {txt(im['PPR_Situação'])}. "
            f"Prevalece, para fins desta análise, a informação de campo. Recomenda-se a atualização "
            f"do registro na base cartográfica.\n\n" + ressalva_area
        )
    else:
        r["identificacao"] = (
            "Os dados disponíveis nos sistemas corporativos e cartoriais são convergentes quanto "
            "ao endereço, à titularidade e à individualização do imóvel.\n\n" + ressalva_area
        )

    r["insercao"] = (
        f"O imóvel integra o projeto urbanístico {txt(im['IMU_PLAN_L'])}, na Região Administrativa "
        f"de {txt(im['REGADMIN'])}, e está sujeito à norma {txt(im['PLANTA_DEC'])}.\n\n"
        f"[Descrever: processo de aprovação do parcelamento e respectivo decreto; objetivos da "
        f"concepção urbanística; enquadramento no PDOT (macrozona e zona); posição do lote dentro "
        f"do conjunto e relação com as vias estruturantes e com as áreas vizinhas.]"
    )

    usos = usos_admitidos(im)
    r["uso"] = (
        f"O imóvel está enquadrado na categoria {txt(im.get('PPR_Destinacao'))}, que admite, "
        f"isolada ou conjuntamente: {usos.lower() if usos else '[descrever os usos admitidos]'}.\n\n"
        f"Os parâmetros do quadro acima foram apurados em: {par['fonte']}.\n\n"
        f"[Avaliar: portes de empreendimento compatíveis com o potencial construtivo; condições "
        f"de acesso e dependência do sistema viário previsto no parcelamento; intervenções viárias "
        f"exigidas pelos órgãos competentes.]"
    )

    sit = txt(im["SITUACAO_VIGENTE"])
    ppr_obs = txt(im.get("PPR_obs:"))
    corpo = (
        f"A titularidade do imóvel está registrada em nome da Terracap. A vistoria da triagem PPR "
        f"classificou o lote como {sit}."
    )
    if ppr_obs != "—":
        corpo += f' Registro da vistoria: "{ppr_obs}"'
    if sit in ("OCUPADO", "CERCADO", "OBSTRUIDO", "CERCADO E VAGO", "CONSTRUIDO"):
        corpo += (
            "\n\nNão foram identificadas informações conclusivas sobre: identidade do ocupante; "
            "extensão exata da área utilizada; origem e duração da ocupação; existência de "
            "instrumento autorizativo; eventual processo de regularização ou desocupação; possível "
            "enquadramento para exercício de direito de preferência.\n\n"
            "A situação dominial e a situação possessória apresentam, portanto, condições "
            "distintas: o imóvel possui titularidade cartorial regular em nome da Terracap, mas "
            "apresenta ocupação física ainda não qualificada nos elementos consultados."
        )
    r["fisica"] = corpo

    gatilhos = {
        "alta tensão": "linha de transmissão de alta tensão e respectiva faixa non aedificandi",
        "transmissão": "linha de transmissão e respectiva faixa de servidão",
        "cerca": "cercamento de terceiros sobre parte da área",
        "invas": "ocupação irregular",
        "entulho": "descarte irregular de resíduos",
        "lixo": "descarte irregular de resíduos",
        "erosão": "processo erosivo",
        "declive": "acidente topográfico relevante",
        "córrego": "curso d'água nas proximidades",
        "app": "área de preservação permanente nas proximidades",
        "poste": "rede de distribuição de energia sobre a área",
        "torre": "estrutura de transmissão implantada sobre a área",
        "antena": "estação de telecomunicações implantada sobre a área",
        "via": "interferência do sistema viário",
    }
    achados = sorted({d for k, d in gatilhos.items() if k in ppr_obs.lower()})
    if achados:
        r["condicionantes"] = (
            "A vistoria e os elementos consultados indicam a presença de: "
            + "; ".join(achados)
            + ".\n\n[Detalhar para cada interferência: extensão da área afetada em m² e em "
            "percentual do lote; existência de solução técnica identificada; custo preliminar "
            "estimado; processo administrativo correspondente; efeito sobre o potencial "
            "construtivo remanescente.]"
        )
    else:
        r["condicionantes"] = (
            "Não foram identificadas, nos elementos consultados, interferências físicas incidentes "
            "sobre o imóvel.\n\n[Confirmar em imagem aérea recente e junto às concessionárias.]"
        )

    r["ambiental"] = (
        "[Registrar: processo ambiental do parcelamento e licença vigente; incidência ou não de "
        "Área de Preservação Permanente sobre o lote; condicionantes de compensação, drenagem e "
        "proteção de cursos d'água; infraestrutura prevista no projeto urbanístico e estágio "
        "efetivo de implantação, incluindo licitações e contratos de obra.]"
    )

    v, criterio = vizinhos(im, cons)
    com_aval = v["TB55_VL_AVALIACAO_RECENTE"].notna().sum()
    com_alien = int((v["TB55_QT_ALIENACOES"].fillna(0) > 0).sum())
    r["mercado"] = (
        f"O quadro acima reúne {len(v)} imóveis da carteira selecionados pelo critério de "
        f"{criterio}. Desses, {com_aval} possuem laudo de avaliação registrado na TB55 e "
        f"{com_alien} apresentam alienação no histórico.\n\n"
        f"[Complementar com evidências externas à carteira: resultados de certames recentes na "
        f"região, valores mínimos e propostas vencedoras, ágios apurados, número de concorrentes, "
        f"empreendimentos em implantação no entorno e efeitos sobre a demanda.]"
    )

    atributos, condicionantes = diagnostico(im, par, v)
    r["avaliacao"] = (
        "Atributos identificados:\n"
        + "\n".join(f"- {a}" for a in atributos)
        + "\n\nCondicionantes identificadas:\n"
        + "\n".join(f"- {c}" for c in condicionantes)
    )

    r["critica"] = (
        "[Juízo do analista: o imóvel reúne ou não características suficientes para ser "
        "reconhecido como ativo com potencial para o desenvolvimento de novos negócios; peso "
        "relativo de cada condicionante; o que efetivamente explica a ausência de comercialização "
        "anterior; o que precisa ser decidido antes da etapa seguinte.]"
    )
    r["encaminhamentos"] = ""
    return r


def diagnostico(im, par, v) -> tuple[list[str], list[str]]:
    """Atributos e condicionantes deduzidos dos dados, para a seção 9."""
    a, c = [], []
    a.append("Unidade imobiliária urbana individualizada e registrada em nome da Terracap.")
    if pd.notna(im["AREA_TERRENO"]):
        a.append(f"Área de terreno de {br_num(im['AREA_TERRENO'], 2)} m².")
    if par["area_maxima"] and pd.notna(par["area_maxima"]):
        a.append(f"Potencial construtivo de {br_num(par['area_maxima'], 2)} m².")
    usos = usos_admitidos(im)
    if usos:
        a.append(f"Usos admitidos: {usos.lower()}.")
    if str(im.get("PPR_Posiçao do imov", "")).upper().startswith("ESQUINA"):
        a.append("Lote de esquina, com mais de uma frente para o sistema viário.")
    if len(v) and v["TB55_QT_ALIENACOES"].fillna(0).sum() > 0:
        a.append("Há registro de alienação em imóveis comparáveis do mesmo conjunto.")
    if im["SITUACAO_VIGENTE"] == "VAGO":
        a.append("Lote vago na última vistoria, sem ocupação aparente.")

    if im["SITUACAO_VIGENTE"] in ("OCUPADO", "CERCADO", "OBSTRUIDO", "CERCADO E VAGO", "CONSTRUIDO"):
        c.append(f"Situação física do lote registrada como {txt(im['SITUACAO_VIGENTE'])} na vistoria.")
    if im["SITUACAO_DIVERGE"]:
        c.append("Divergência entre a situação registrada na base cartográfica e a apurada em campo.")
    if not im["TEM_AVALIACAO"]:
        c.append("Imóvel sem laudo de avaliação registrado na TB55.")
    if im.get("EM_TB55") == "Não":
        c.append("Imóvel sem histórico cartorial e processual vinculado na TB55.")
    if not par["vinculado_luos"] and str(im.get("LUOS_MATCH_STATUS", "")).startswith(
        ("Área do lote", "Ambíguo", "UOS")
    ):
        c.append(f"Vinculação à base LUOS pendente: {txt(im['LUOS_MATCH_STATUS'])}.")
    if im["SEM_POTENCIAL"]:
        c.append("Coeficiente de aproveitamento nulo ou não informado na base cartográfica.")
    if not c:
        c.append("Não foram identificadas condicionantes nos elementos consultados.")
    return a, c


def encaminhamentos_sugeridos(im, par) -> list[str]:
    """Checklist adaptado ao imóvel, somado aos encaminhamentos padrão."""
    s = []
    if im["SITUACAO_VIGENTE"] in ("OCUPADO", "CERCADO", "OBSTRUIDO", "CERCADO E VAGO", "CONSTRUIDO"):
        s.append(
            "Qualificar a ocupação existente: identidade do ocupante, extensão da área utilizada, "
            "origem e duração, instrumento autorizativo e eventual direito de preferência."
        )
    if not im["TEM_AVALIACAO"]:
        s.append("Solicitar avaliação do imóvel à área competente, para subsidiar a decisão comercial.")
    if im["SITUACAO_DIVERGE"]:
        s.append("Atualizar a situação do lote na base cartográfica conforme a vistoria de campo.")
    if not par["vinculado_luos"] and str(im.get("LUOS_MATCH_STATUS", "")).startswith(
        ("Área do lote", "Ambíguo", "UOS")
    ):
        s.append("Concluir a vinculação do lote à base LUOS revisada, resolvendo os candidatos em aberto.")
    return s + ENCAMINHAMENTOS_PADRAO


# ----------------------------------------------------------------------------
# Saídas
# ----------------------------------------------------------------------------
def montar_markdown(im, textos, quadros, marcados, autor) -> str:
    p = [
        "# PPR 2026 — Prospecção de Imóveis Urbanos",
        "",
        f"## Imóvel GIU {txt(im['CD_IMOVEL'])}",
        f"**{txt(im['ENDERECO'])} — {txt(im['REGADMIN'])}**",
        "",
        f"*Justificativa Técnica Preliminar · {data_extenso()}"
        + (f" · {autor}*" if autor else "*"),
        "",
    ]
    for chave, num, titulo, tem_quadro, _ in SECOES:
        p += [f"### {num}. {titulo}", ""]
        if tem_quadro and chave in quadros:
            q = quadros[chave]
            if q and isinstance(q[0], list) and len(q[0]) > 2:
                p.append("| " + " | ".join(q[0]) + " |")
                p.append("|" + "---|" * len(q[0]))
                for linha in q[1:]:
                    p.append("| " + " | ".join(str(x).replace("|", "/") for x in linha) + " |")
            else:
                p += ["| Informação | Descrição |", "|---|---|"]
                for rot, val in q:
                    p.append(f"| {rot} | {str(val).replace('|', '/')} |")
            p.append("")
        if chave == "encaminhamentos" and marcados:
            for i, e in enumerate(marcados, 1):
                p.append(f"{i}. {e}")
            p.append("")
        t = (textos.get(chave) or "").strip()
        if t:
            p += [t, ""]
    p += ["---", "", "Documento gerado pelo Painel de Consulta de Imóveis · Terracap."]
    return "\n".join(p)


def montar_docx(im, textos, quadros, marcados, autor) -> bytes:
    """Documento Word com a identidade visual institucional."""
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt, RGBColor

    verde = RGBColor(0x00, 0x6D, 0x33)
    azul = RGBColor(0x00, 0x41, 0x6D)

    doc = Document()
    for s in doc.sections:
        s.top_margin = s.bottom_margin = Cm(2.5)
        s.left_margin = s.right_margin = Cm(3.0)
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    def paragrafo(texto, tamanho=11, negrito=False, cor=None, alinhar=None,
                  espaco_antes=0, espaco_depois=6, recuo=None):
        p = doc.add_paragraph()
        if alinhar:
            p.alignment = alinhar
        if recuo:
            p.paragraph_format.first_line_indent = Cm(recuo)
        p.paragraph_format.space_before = Pt(espaco_antes)
        p.paragraph_format.space_after = Pt(espaco_depois)
        run = p.add_run(texto)
        run.bold = negrito
        run.font.size = Pt(tamanho)
        if cor:
            run.font.color.rgb = cor
        return p

    def sombrear(celula, cor_hex):
        tc = celula._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), cor_hex)
        tc.append(shd)

    # Cabeçalho
    paragrafo("PPR 2026 — PROSPECÇÃO DE IMÓVEIS URBANOS", 13, True, verde,
              WD_ALIGN_PARAGRAPH.CENTER, espaco_depois=2)
    paragrafo("Justificativa Técnica Preliminar", 10, False, azul,
              WD_ALIGN_PARAGRAPH.CENTER, espaco_depois=14)
    paragrafo(f"Imóvel GIU {txt(im['CD_IMOVEL'])}", 12, True, cor=azul, espaco_depois=1)
    paragrafo(txt(im["ENDERECO"]), 11, True, espaco_depois=1)
    paragrafo(f"{txt(im['REGADMIN'])} · Condição {txt(im['CD_DS_COND'])}", 10, cor=RGBColor(0x5A, 0x6B, 0x7B))
    rodape_topo = data_extenso() + (f" · {autor}" if autor else "")
    paragrafo(rodape_topo, 9, cor=RGBColor(0x5A, 0x6B, 0x7B), espaco_depois=14)

    def inserir_quadro(q):
        if q and isinstance(q[0], list) and len(q[0]) > 2:
            cabecalho, corpo = q[0], q[1:]
        else:
            cabecalho, corpo = ["Informação", "Descrição"], q
        t = doc.add_table(rows=1, cols=len(cabecalho))
        t.style = "Table Grid"
        for i, c in enumerate(cabecalho):
            cel = t.rows[0].cells[i]
            cel.text = ""
            run = cel.paragraphs[0].add_run(str(c))
            run.bold = True
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            sombrear(cel, "006D33")
        for j, linha in enumerate(corpo):
            cels = t.add_row().cells
            for i, val in enumerate(linha[: len(cabecalho)]):
                cels[i].text = ""
                run = cels[i].paragraphs[0].add_run(str(val))
                run.font.size = Pt(9.5)
                if i == 0 and len(cabecalho) == 2:
                    run.bold = True
            if j % 2 == 1:
                for cel in cels:
                    sombrear(cel, "F2F6F3")
        doc.add_paragraph().paragraph_format.space_after = Pt(4)

    for chave, num, titulo, tem_quadro, _ in SECOES:
        paragrafo(f"{num}. {titulo}", 11.5, True, verde, espaco_antes=12, espaco_depois=6)
        if tem_quadro and chave in quadros:
            inserir_quadro(quadros[chave])
        if chave == "encaminhamentos" and marcados:
            for e in marcados:
                p = doc.add_paragraph(e, style="List Number")
                p.paragraph_format.space_after = Pt(4)
                for run in p.runs:
                    run.font.size = Pt(11)
        for bloco in (textos.get(chave) or "").split("\n"):
            bloco = bloco.strip()
            if not bloco:
                continue
            if bloco.startswith(("- ", "• ")):
                p = doc.add_paragraph(bloco[2:], style="List Bullet")
                p.paragraph_format.space_after = Pt(3)
                for run in p.runs:
                    run.font.size = Pt(11)
            else:
                p = paragrafo(bloco, 11, recuo=1.25, espaco_depois=6)
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    paragrafo(
        "Documento gerado pelo Painel de Consulta de Imóveis a partir da carteira consolidada "
        "(base cartográfica, triagem PPR, TB55 e base LUOS revisada). Os campos preenchidos "
        "automaticamente devem ser conferidos antes da tramitação.",
        8, cor=RGBColor(0x5A, 0x6B, 0x7B), espaco_antes=18,
    )

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
